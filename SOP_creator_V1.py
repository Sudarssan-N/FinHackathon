from pynput import keyboard, mouse
import speech_recognition as sr
import pyautogui
from docx import Document
from datetime import datetime
from PIL import Image, ImageDraw

text_entries = []
image_paths = []

def capture_screenshot_with_highlight(x, y):
    # Capture a screenshot
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    screenshot_path = f"screenshot_{timestamp}.png"
    screenshot = pyautogui.screenshot()

    # Convert screenshot to a PIL Image to draw on it
    screenshot = screenshot.convert("RGB")
    draw = ImageDraw.Draw(screenshot)

    # Define the size of the rectangle around the cursor
    rect_size = 50
    left = x - rect_size // 2
    top = y - rect_size // 2
    right = x + rect_size // 2
    bottom = y + rect_size // 2

    # Draw a rectangle around the cursor position
    draw.rectangle([left, top, right, bottom], outline="red", width=3)

    # Save the screenshot with the highlighted cursor area
    highlighted_path = f"highlighted_{timestamp}.png"
    screenshot.save(highlighted_path)
    print(f"Screenshot with highlight saved as {highlighted_path}")
    return highlighted_path

def transcribe_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening for speech...")
        audio = recognizer.listen(source)
        try:
            # Recognize speech using Google Web Speech API
            text = recognizer.recognize_google(audio)
            print(f"Transcribed Text: {text}")
            return text
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand the audio")
            return ""
        except sr.RequestError as e:
            print(f"Could not request results from Google Speech Recognition service; {e}")
            return ""

def create_document(text_entries, image_paths):
    doc = Document()
    doc.add_heading("SOP Document", level=1)
    
    for text in text_entries:
        doc.add_paragraph(text)
    
    for image_path in image_paths:
        doc.add_picture(image_path)
    
    doc_filename = f"SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(doc_filename)
    print(f"Document saved as {doc_filename}")

def on_click(x, y, button, pressed):
    if pressed:
        print(f"Mouse clicked at ({x}, {y})")
        image_path = capture_screenshot_with_highlight(x, y)
        image_paths.append(image_path)
        return False  # Stop listener after the first click to avoid multiple captures

def on_press(key):
    try:
        if key == keyboard.HotKey.parse('<ctrl>+<shift>+f'):
            text = transcribe_speech()
            if text:
                text_entries.append(text)
        elif key == keyboard.HotKey.parse('<ctrl>+<shift>+p'):
            print("Click anywhere on the screen to capture a screenshot with a highlight.")
            with mouse.Listener(on_click=on_click) as listener:
                listener.join()
        elif key == keyboard.HotKey.parse('<ctrl>+<shift>+q'):
            create_document(text_entries, image_paths)
            return False  # Stop listener
    except AttributeError:
        pass

if __name__ == "__main__":
    print("Press 'ctrl+shift+f' to start/stop speech recognition.")
    print("Press 'ctrl+shift+p' to capture a screenshot.")
    print("Press 'ctrl+shift+q' to create the document and exit.")
    
    with keyboard.Listener(on_press=on_press) as listener:
        listener.join()


let eventSource = null;
let fetchController = null;

// Function to start the subscription with custom headers
function startSSE() {
    if (eventSource !== null) {
        console.log("SSE connection is already open.");
        return;
    }

    fetchController = new AbortController();
    const signal = fetchController.signal;

    // Make a Fetch request with custom headers
    fetch('http://localhost:8080/subscribe', {
        method: 'GET',
        headers: {
            'Authorization': 'Bearer YOUR_TOKEN_HERE',  // Add your custom headers here
            'Custom-Header': 'YourCustomValue'
        },
        signal: signal
    }).then(response => {
        if (!response.ok) {
            throw new Error(`HTTP error! status: ${response.status}`);
        }
        return response.body.getReader();
    }).then(reader => {
        const decoder = new TextDecoder();
        processStream(reader, decoder);
    }).catch(error => {
        console.error("SSE connection error:", error);
        addMessage("SSE connection error. Check the console for more details.");
        stopSSE(); // Stop the connection in case of error
    });

    console.log("SSE connection opened using Fetch API.");
    addMessage("SSE connection opened using Fetch API.");
}

// Function to process the stream from the Fetch API
function processStream(reader, decoder) {
    reader.read().then(function processText({ done, value }) {
        if (done) {
            console.log("SSE connection closed by the server.");
            addMessage("SSE connection closed by the server.");
            stopSSE();
            return;
        }

        // Decode and handle the chunk of data
        const text = decoder.decode(value, { stream: true });
        console.log("New message received:", text);
        addMessage(text);

        // Continue reading the stream
        return reader.read().then(processText);
    }).catch(error => {
        console.error("Error while reading stream:", error);
        addMessage("Error while reading stream.");
        stopSSE();
    });
}

// Function to stop the SSE connection
function stopSSE() {
    if (fetchController !== null) {
        fetchController.abort();
        fetchController = null;
        console.log("SSE connection closed.");
        addMessage("SSE connection closed.");
    }
}

// Function to add a message to the messages div
function addMessage(message) {
    const messagesDiv = document.getElementById("messages");
    const newMessage = document.createElement("p");
    newMessage.textContent = message;
    messagesDiv.appendChild(newMessage);
    messagesDiv.scrollTop = messagesDiv.scrollHeight; // Scroll to the bottom
}