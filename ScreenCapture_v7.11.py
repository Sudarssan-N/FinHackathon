import tkinter as tk
from tkinter import messagebox
import pyautogui
import keyboard
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import getpass

# Global variables to keep track of steps and the document
steps = []
doc = None

# Function to take a screenshot with the cursor highlighted
def take_screenshot():
    # Get the current mouse position
    mouse_x, mouse_y = pyautogui.position()
    
    # Take the screenshot
    screenshot = pyautogui.screenshot()
    
    # Convert the screenshot to an editable format using Pillow (PIL)
    screenshot = screenshot.convert("RGBA")
    
    # Create a drawing object to draw the cursor
    draw = ImageDraw.Draw(screenshot, "RGBA")
    
    # Define the rectangle dimensions and color
    rectangle_width = 80
    rectangle_height = 40
    rectangle_color = (255, 0, 0, 255)  # Red
    
    # Calculate the top-left and bottom-right coordinates of the rectangle
    top_left = (mouse_x - rectangle_width // 2, mouse_y - rectangle_height // 2)
    bottom_right = (mouse_x + rectangle_width // 2, mouse_y + rectangle_height // 2)
    
    # Draw the rectangle border around the cursor
    draw.rectangle([top_left, bottom_right], outline=rectangle_color, width=3)
    
    # Convert back to RGB before saving
    screenshot = screenshot.convert("RGB")
    
    # Save the screenshot with a timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_path = f"screenshot_{timestamp}.png"
    screenshot.save(file_path)

    return file_path

# Function to update the displayed steps in the GUI
def update_step_display():
    for widget in steps_frame.winfo_children():
        widget.destroy()  # Clear the current display
    
    for i, step in enumerate(steps):
        step_label = tk.Label(steps_frame, text=f"Step {i+1}: {step['description']}")
        step_label.grid(row=i, column=0, padx=10, pady=5)
        
        # Add + button to insert a step
        add_button = tk.Button(steps_frame, text="+", command=lambda index=i: insert_step(index))
        add_button.grid(row=i, column=1, padx=5)
        
        # Add Edit button to edit the step
        edit_button = tk.Button(steps_frame, text="Edit", command=lambda index=i: edit_step(index))
        edit_button.grid(row=i, column=2, padx=5)

# Function to add a new step between existing steps
def insert_step(index):
    def on_print_screen():
        new_step_description = description_text.get("1.0", tk.END).strip()
        if not new_step_description:
            messagebox.showerror("Error", "Please enter a description for the new step.")
            return

        # Take a new screenshot for the step
        file_path = take_screenshot()
        
        # Insert the new step into the list of steps
        steps.insert(index + 1, {"description": new_step_description, "screenshot": file_path})
        
        # Update the step display and clear the description text box
        update_step_display()
        description_text.delete("1.0", tk.END)
    
    keyboard.wait("ctrl + Shift + P")  # Wait for the Print Screen key press
    on_print_screen()

# Function to edit an existing step
def edit_step(index):
    def save_edited_step():
        new_description = edit_description_text.get("1.0", tk.END).strip()
        if not new_description:
            messagebox.showerror("Error", "Description cannot be empty.")
            return
        
        # Update the description and screenshot for the step
        steps[index]["description"] = new_description
        steps[index]["screenshot"] = take_screenshot()  # Take a new screenshot
        
        # Close the edit window and update the step display
        edit_window.destroy()
        update_step_display()
    
    # Open a new window for editing the step
    edit_window = tk.Toplevel(root)
    edit_window.title(f"Edit Step {index+1}")
    
    # Show the current description
    edit_description_label = tk.Label(edit_window, text="Edit Description:")
    edit_description_label.pack(pady=5)
    
    edit_description_text = tk.Text(edit_window, width=40, height=5)
    edit_description_text.insert(tk.END, steps[index]["description"])
    edit_description_text.pack(pady=5)
    
    # Button to save the edited step
    save_button = tk.Button(edit_window, text="Save", command=save_edited_step)
    save_button.pack(pady=10)

# Function to save the document
def save_document():
    if not steps:
        messagebox.showerror("Error", "No steps to save.")
        return
    
    global doc
    if doc is None:
        doc = Document()
        header_text = header_entry.get()
        doc.add_heading(header_text, level=0).alignment = 1  # Centered title
        
        # Add metadata
        today_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        user_id = getpass.getuser()
        doc.add_paragraph(f"Date: {today_date}").alignment = 1  # Centered date
        doc.add_paragraph(f"User ID: {user_id}").alignment = 1  # Centered user ID
        doc.add_page_break()
    
    for i, step in enumerate(steps):
        doc.add_heading(f"Step {i+1}", level=1)
        doc.add_paragraph(step['description'])
        doc.add_picture(step['screenshot'], width=Inches(6))
    
    file_name = file_name_entry.get().strip()
    if file_name:
        doc.save(f"{file_name}.docx")
        messagebox.showinfo("Success", f"Document saved as {file_name}.docx")
    else:
        messagebox.showerror("Error", "Please provide a valid file name.")

# Create the main window
root = tk.Tk()
root.title("Screenshot Documenter with Step Management")

# GUI elements
tk.Label(root, text="Document Header Name:").pack(pady=5)
header_entry = tk.Entry(root, width=40)
header_entry.pack(pady=5)

tk.Label(root, text="Step Description:").pack(pady=5)
description_text = tk.Text(root, width=40, height=5)
description_text.pack(pady=5)

# Frame to display steps dynamically
steps_frame = tk.Frame(root)
steps_frame.pack(pady=10)

# Buttons to start and save document
tk.Button(root, text="Add Step", command=lambda: insert_step(len(steps) - 1)).pack(pady=5)
tk.Label(root, text="File Name to Save Document:").pack(pady=5)
file_name_entry = tk.Entry(root, width=40)
file_name_entry.pack(pady=5)
tk.Button(root, text="Save Document", command=save_document).pack(pady=10)

# Start the GUI event loop
root.mainloop()
