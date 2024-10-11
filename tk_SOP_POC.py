import tkinter as tk
import pyautogui
import speech_recognition as sr
from PIL import Image, ImageDraw
from datetime import datetime
import os
from docx import Document

# Get the Downloads folder path
downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

text_entries = []
image_paths = []

def capture_screenshot():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    screenshot_path = os.path.join(downloads_folder, f"screenshot_{timestamp}.png")
    screenshot = pyautogui.screenshot()
    cursor_x, cursor_y = pyautogui.position()
    highlight_rectangle(screenshot, cursor_x, cursor_y, screenshot_path)
    log_message(f"Screenshot saved at: {screenshot_path}")
    return screenshot_path

def highlight_rectangle(image, x, y, save_path):
    image = image.convert("RGB")
    draw = ImageDraw.Draw(image)
    rect_width, rect_height = 100, 50
    top_left = (x - rect_width // 2, y - rect_height // 2)
    bottom_right = (x + rect_width // 2, y + rect_height // 2)
    draw.rectangle([top_left, bottom_right], outline="red", width=3)
    image.save(save_path)

def transcribe_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        log_message("Listening for speech...")
        audio = recognizer.listen(source)
        try:
            text = recognizer.recognize_google(audio)
            text_entries.append(text)
            log_message(f"Transcribed Text: {text}")
        except sr.UnknownValueError:
            log_message("Could not understand audio")
        except sr.RequestError as e:
            log_message(f"Request error: {e}")

def create_document():
    doc = Document()
    doc.add_heading("SOP Document", level=1)
    
    for text in text_entries:
        doc.add_paragraph(text)
    
    for image_path in image_paths:
        doc.add_picture(image_path)
    
    doc_filename = os.path.join(downloads_folder, f"SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_filename)
    log_message(f"Document saved at: {doc_filename}")

def log_message(message):
    log_text.insert(tk.END, message + "\n")
    log_text.yview(tk.END)

# Create the main window
root = tk.Tk()
root.title("SOP Automation")
root.geometry("300x400")
root.attributes("-topmost", True)  # Keep the window on top of others

# Create buttons
btn_transcribe = tk.Button(root, text="Start Transcription", command=transcribe_speech)
btn_transcribe.pack(pady=10)

btn_screenshot = tk.Button(root, text="Capture Screenshot", command=lambda: image_paths.append(capture_screenshot()))
btn_screenshot.pack(pady=10)

btn_create_doc = tk.Button(root, text="Create Document", command=create_document)
btn_create_doc.pack(pady=10)

# Create a text widget for logs
log_text = tk.Text(root, height=10, width=35)
log_text.pack(pady=10)

# Start the application
root.mainloop()
