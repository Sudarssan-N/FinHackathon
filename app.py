from flask import Flask, render_template, request, jsonify
import speech_recognition as sr
import pyautogui
from datetime import datetime
from PIL import Image, ImageDraw
import os

app = Flask(__name__)

# Get the Downloads folder path
downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

text_entries = []
image_paths = []

def capture_screenshot():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    screenshot_path = os.path.join(downloads_folder, f"screenshot_{timestamp}.png")
    screenshot = pyautogui.screenshot()
    cursor_x, cursor_y = pyautogui.position()
    highlight_rectangle(screenshot, cursor_x, cursor_y, screenshot_path)
    return screenshot_path

def highlight_rectangle(image, x, y, save_path):
    image = image.convert("RGB")
    draw = ImageDraw.Draw(image)
    rect_width, rect_height = 100, 50
    top_left = (x - rect_width // 2, y - rect_height // 2)
    bottom_right = (x + rect_width // 2, y + rect_height // 2)
    draw.rectangle([top_left, bottom_right], outline="red", width=3)
    image.save(save_path)

def transcribe_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        audio = recognizer.listen(source)
        try:
            text = recognizer.recognize_google(audio)
            text_entries.append(text)
            return text
        except sr.UnknownValueError:
            return "Could not understand audio"
        except sr.RequestError as e:
            return f"Request error: {e}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/transcribe', methods=['POST'])
def transcribe():
    text = transcribe_speech()
    return jsonify({"text": text})

@app.route('/screenshot', methods=['POST'])
def screenshot():
    image_path = capture_screenshot()
    image_paths.append(image_path)
    return jsonify({"image_path": image_path})

@app.route('/create_document', methods=['POST'])
def create_document():
    from docx import Document
    doc = Document()
    doc.add_heading("SOP Document", level=1)
    for text in text_entries:
        doc.add_paragraph(text)
    for image_path in image_paths:
        doc.add_picture(image_path)
    
    doc_filename = os.path.join(downloads_folder, f"SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_filename)
    return jsonify({"document_path": doc_filename})

if __name__ == "__main__":
    app.run(debug=True)
