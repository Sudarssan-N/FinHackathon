import tkinter as tk
from tkinter import messagebox
import pyautogui
import keyboard
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageTk
import getpass
import threading

# Global variables
steps = []
doc = None
is_capturing = False
capture_interval = 1000
current_screenshots = []  # To store multiple screenshots for a single step
is_recording = False  # Recording flag for the independent recording feature

# Function to take a screenshot and highlight the cursor position
def take_screenshot():
    mouse_x, mouse_y = pyautogui.position()
    screenshot = pyautogui.screenshot()
    screenshot = screenshot.convert("RGBA")
    draw = ImageDraw.Draw(screenshot, "RGBA")
    rectangle_width = 80
    rectangle_height = 40
    rectangle_color = (255, 0, 0, 255)
    top_left = (mouse_x - rectangle_width // 2, mouse_y - rectangle_height // 2)
    bottom_right = (mouse_x + rectangle_width // 2, mouse_y + rectangle_height // 2)
    draw.rectangle([top_left, bottom_right], outline=rectangle_color, width=3)
    screenshot = screenshot.convert("RGB")
    now = datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")
    file_path = f"screenshot_{timestamp}.png"
    screenshot.save(file_path)
    return file_path

# Function to create the Word document
def save_to_word():
    global doc
    doc = Document()
    header_text = header_entry.get()
    user_id = getpass.getuser()
    today_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_heading(header_text, level=0).alignment = 1
    doc.add_paragraph(f"Date: {today_date}", style='Normal').alignment = 1
    doc.add_paragraph(header_description.get(), style='Normal').alignment = 1
    doc.add_paragraph(f"User ID: {user_id}", style='Normal').alignment = 1
    doc.add_page_break()

    for step in steps:
        doc.add_heading(f"Step {step['step_number']}", level=1)
        doc.add_paragraph(step['description'])
        for screenshot_path in step['screenshot_paths']:
            doc.add_picture(screenshot_path, width=Inches(6))

# Function to add a step at a specific position
def add_step(step_number=None, with_screenshots=True):
    image_description = description_text.get("1.0", tk.END).strip()
    if not image_description:
        messagebox.showerror("Error", "Please enter a description for the step.")
        return

    screenshot_paths = current_screenshots if with_screenshots else []

    if step_number is None:
        step_number = len(steps) + 1
    
    new_step = {
        'step_number': step_number,
        'description': image_description,
        'screenshot_paths': screenshot_paths
    }

    # Insert step and renumber subsequent steps
    steps.insert(step_number - 1, new_step)
    for i, step in enumerate(steps):
        step['step_number'] = i + 1

    update_gui()
    description_text.delete("1.0", tk.END)
    clear_screenshot_cache()

# Function to start capturing multiple screenshots in a separate thread
def start_screenshot_capture():
    global is_capturing
    is_capturing = True
    messagebox.showinfo("Info", "Screenshot capturing started. Press 'Print Screen' to capture screenshots, and click 'Stop Screenshot' to finish.")
    
    # Start the thread for capturing screenshots based on key presses
    threading.Thread(target=monitor_screenshot_key).start()

# Monitor for 'Print Screen' keypress for screenshot capture
def monitor_screenshot_key():
    global is_capturing
    while is_capturing:
        if keyboard.is_pressed('print screen'):
            screenshot_path = take_screenshot()
            current_screenshots.append(screenshot_path)
            messagebox.showinfo("Success", f"Screenshot captured: {screenshot_path}")
            keyboard.wait('print screen')  # To avoid multiple captures on one press

# Function to stop capturing screenshots
def stop_screenshot_capture():
    global is_capturing
    is_capturing = False
    messagebox.showinfo("Info", "Screenshot capturing stopped.")

# Function to start recording (independent of Word document)
def start_recording():
    global is_recording
    is_recording = True
    messagebox.showinfo("Info", "Recording started.")

# Function to stop recording (independent of Word document)
def stop_recording():
    global is_recording
    is_recording = False
    messagebox.showinfo("Info", "Recording stopped.")

# Function to clear the screenshot cache after a step is added
def clear_screenshot_cache():
    global current_screenshots
    current_screenshots = []

# Function to edit an existing step
def edit_step(step_number):
    for step in steps:
        if step['step_number'] == step_number:
            description_text.delete("1.0", tk.END)
            description_text.insert(tk.END, step['description'])
            screenshot_paths = current_screenshots if is_capturing else step['screenshot_paths']
            step['description'] = description_text.get("1.0", tk.END).strip()
            step['screenshot_paths'] = screenshot_paths
            update_gui()
            break

# Function to update the GUI with steps
def update_gui():
    for widget in step_frame.winfo_children():
        widget.destroy()

    for step in steps:
        step_label = tk.Label(step_frame, text=f"Step {step['step_number']}: {step['description'][:30]}...")
        step_label.pack(anchor='w')

        edit_button = tk.Button(step_frame, text="Edit", command=lambda s=step['step_number']: edit_step(s))
        edit_button.pack(anchor='w')

        add_button = tk.Button(step_frame, text="+", command=lambda s=step['step_number']: add_step(s))
        add_button.pack(anchor='w')

# Function to save the document
def save_document():
    if not steps:
        messagebox.showerror("Error", "No steps to save.")
        return

    save_to_word()
    file_name = file_name_entry.get()
    if file_name:
        doc.save(f"{file_name}.docx")
        messagebox.showinfo("Success", f"Document saved as {file_name}.docx")
    else:
        messagebox.showerror("Error", "Please enter a valid file name.")

# Create the main window
root = tk.Tk()
root.title("Screenshot Documenter")

# Load and display the logo in the top left corner
logo = Image.open("OIP.jpeg")  # Replace with the path to your logo
logo = logo.resize((100, 100), Image.Resampling.LANCZOS)  # Resize the logo if necessary
logo_tk = ImageTk.PhotoImage(logo)

logo_label = tk.Label(root, image=logo_tk)
logo_label.place(x=10, y=10)  # Adjust the position of the logo

# Create GUI elements
tk.Label(root, text="Document Header Name:").pack(pady=5)
header_entry = tk.Entry(root, width=40)
header_entry.pack(pady=5)

tk.Label(root, text="Document Description:").pack(pady=5)
header_description = tk.Entry(root, width=40)
header_description.pack(pady=5)

tk.Label(root, text="Step Description:").pack(pady=5)
description_text = tk.Text(root, width=40, height=5)
description_text.pack(pady=5)

# Screenshot capture buttons
tk.Button(root, text="Start Screenshot Capture", command=start_screenshot_capture).pack(pady=10)
tk.Button(root, text="Stop Screenshot Capture", command=stop_screenshot_capture).pack(pady=10)

# Add step buttons
tk.Button(root, text="Add New Step (Without Screenshot)", command=lambda: add_step(with_screenshots=False)).pack(pady=10)
tk.Button(root, text="Add New Step (With Screenshots)", command=lambda: add_step(with_screenshots=True)).pack(pady=10)

# Recording buttons (independent of Word document)
tk.Button(root, text="Start Recording", command=start_recording).pack(pady=10)
tk.Button(root, text="Stop Recording", command=stop_recording).pack(pady=10)

step_frame = tk.Frame(root)
step_frame.pack(pady=10)

tk.Label(root, text="File Name to Save Document:").pack(pady=5)
file_name_entry = tk.Entry(root, width=40)
file_name_entry.pack(pady=5)

tk.Button(root, text="Save Document", command=save_document).pack(pady=10)

# Start the GUI event loop
root.mainloop()
