import speech_recognition as sr
import pyautogui
import keyboard
from docx import Document
from datetime import datetime
from PIL import Image, ImageDraw
import os

# Get the Downloads folder path
downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

def capture_screenshot():
    # Capture screenshot and save it with a timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    screenshot_path = os.path.join(downloads_folder, f"screenshot_{timestamp}.png")
    screenshot = pyautogui.screenshot()
    
    # Get cursor position and draw a rectangle around it
    cursor_x, cursor_y = pyautogui.position()
    highlight_rectangle(screenshot, cursor_x, cursor_y, screenshot_path)
    print(f"Screenshot with cursor highlight saved as {screenshot_path}")
    return screenshot_path

def highlight_rectangle(image, x, y, save_path):
    # Convert screenshot to an editable image
    image = image.convert("RGB")
    draw = ImageDraw.Draw(image)
    
    # Define rectangle dimensions (adjust as needed)
    rect_width, rect_height = 100, 50
    top_left = (x - rect_width // 2, y - rect_height // 2)
    bottom_right = (x + rect_width // 2, y + rect_height // 2)
    
    # Draw a rectangle around the cursor position
    draw.rectangle([top_left, bottom_right], outline="red", width=3)
    image.save(save_path)

def transcribe_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening for speech...")
        audio = recognizer.listen(source)
        try:
            # Recognize speech using Google Web Speech API
            text = recognizer.recognize_google(audio)
            print(f"Transcribed Text: {text}")
            return text
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand the audio")
            return ""
        except sr.RequestError as e:
            print(f"Could not request results from Google Speech Recognition service; {e}")
            return ""

def create_document(text_entries, image_paths):
    doc = Document()
    doc.add_heading("SOP Document", level=1)
    
    for text in text_entries:
        doc.add_paragraph(text)
    
    for image_path in image_paths:
        doc.add_picture(image_path)
    
    doc_filename = os.path.join(downloads_folder, f"SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_filename)
    print(f"Document saved as {doc_filename}")

if __name__ == "__main__":
    text_entries = []
    image_paths = []

    print("Press 'ctrl+shift+s' to start/stop speech recognition.")
    print("Press 'ctrl+shift+p' to capture a screenshot.")
    print("Press 'ctrl+shift+q' to create the document and exit.")

    while True:
        if keyboard.is_pressed('ctrl+shift+s'):
            text = transcribe_speech()
            if text:
                text_entries.append(text)
        
        if keyboard.is_pressed('ctrl+shift+p'):
            image_path = capture_screenshot()
            image_paths.append(image_path)
        
        if keyboard.is_pressed('ctrl+shift+q'):
            create_document(text_entries, image_paths)
            break
